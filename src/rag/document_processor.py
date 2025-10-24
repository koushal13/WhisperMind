"""
Document processing for RAG system.
"""

import os
import asyncio
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import hashlib

# Document processing libraries
try:
    import pypdf
    from docx import Document as DocxDocument
    import markdown
    from bs4 import BeautifulSoup
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Document container for processed text."""
    content: str
    metadata: Dict[str, Any]
    doc_id: str


class DocumentProcessor:
    """Processes various document formats for RAG."""
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.html': 'html',
        '.htm': 'html'
    }
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Size of text chunks for processing
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    async def process_directory(self, directory_path: str) -> List[Document]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of processed documents
        """
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.warning(f"Directory not found: {directory_path}")
            return documents
        
        # Find all supported files
        supported_files = []
        for ext in self.SUPPORTED_EXTENSIONS:
            supported_files.extend(directory.rglob(f"*{ext}"))
        
        logger.info(f"Found {len(supported_files)} supported files")
        
        # Process each file
        for file_path in supported_files:
            try:
                file_documents = await self.process_file(str(file_path))
                documents.extend(file_documents)
                logger.info(f"Processed {file_path.name}: {len(file_documents)} chunks")
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
        
        return documents
    
    async def process_file(self, file_path: str) -> List[Document]:
        """
        Process a single file and return document chunks.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of document chunks
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Extract text based on file type
        if extension == '.txt':
            text = await self._read_text_file(file_path)
        elif extension == '.md':
            text = await self._read_markdown_file(file_path)
        elif extension == '.pdf':
            text = await self._read_pdf_file(file_path)
        elif extension == '.docx':
            text = await self._read_docx_file(file_path)
        elif extension in ['.html', '.htm']:
            text = await self._read_html_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return []
        
        # Create metadata
        metadata = {
            'source': str(path),
            'filename': path.name,
            'extension': extension,
            'size': path.stat().st_size,
            'modified': path.stat().st_mtime,
            'type': self.SUPPORTED_EXTENSIONS[extension]
        }
        
        # Split into chunks
        chunks = self._split_text(text)
        
        # Create documents
        documents = []
        for i, chunk in enumerate(chunks):
            doc_id = self._generate_doc_id(file_path, i)
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'chunk_index': i,
                'chunk_count': len(chunks)
            })
            
            documents.append(Document(
                content=chunk,
                metadata=chunk_metadata,
                doc_id=doc_id
            ))
        
        return documents
    
    async def _read_text_file(self, file_path: str) -> str:
        """Read plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    async def _read_markdown_file(self, file_path: str) -> str:
        """Read markdown file and convert to text."""
        try:
            import markdown
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert markdown to HTML then to text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        except ImportError:
            # Fallback to reading as plain text
            return await self._read_text_file(file_path)
    
    async def _read_pdf_file(self, file_path: str) -> str:
        """Read PDF file."""
        try:
            import pypdf
            text = ""
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except ImportError:
            raise ImportError("pypdf library is required for PDF processing. Install with: pip install pypdf")
    
    async def _read_docx_file(self, file_path: str) -> str:
        """Read DOCX file."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise ImportError("python-docx library is required for DOCX processing. Install with: pip install python-docx")
    
    async def _read_html_file(self, file_path: str) -> str:
        """Read HTML file."""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
        except ImportError:
            raise ImportError("beautifulsoup4 library is required for HTML processing. Install with: pip install beautifulsoup4")
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap.
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                # Last chunk
                chunks.append(text[start:])
                break
            
            # Try to find a good break point (sentence end, paragraph, etc.)
            chunk_text = text[start:end]
            
            # Look for sentence endings
            for delimiter in ['. ', '.\n', '!\n', '?\n']:
                last_delimiter = chunk_text.rfind(delimiter)
                if last_delimiter > self.chunk_size * 0.5:  # Don't make chunks too small
                    end = start + last_delimiter + len(delimiter)
                    break
            
            chunks.append(text[start:end])
            start = end - self.chunk_overlap
        
        return [chunk.strip() for chunk in chunks if chunk.strip()]
    
    def _generate_doc_id(self, file_path: str, chunk_index: int) -> str:
        """Generate unique document ID."""
        content = f"{file_path}_{chunk_index}"
        return hashlib.md5(content.encode()).hexdigest()


# Example usage
if __name__ == "__main__":
    async def main():
        processor = DocumentProcessor()
        
        # Process a directory
        documents = await processor.process_directory("data/documents")
        
        print(f"Processed {len(documents)} document chunks")
        
        if documents:
            print(f"First document preview: {documents[0].content[:200]}...")
            print(f"Metadata: {documents[0].metadata}")
    
    asyncio.run(main())